from fastapi import FastAPI, UploadFile, File, HTTPException, BackgroundTasks, File, Form
from fastapi.responses import FileResponse
import tempfile

from fastapi.middleware.cors import CORSMiddleware
from schemas import TemplateOut, DocumentOut, TemplateVariableOut   
import os
from markdown_it import MarkdownIt
from typing import Dict, List
import json
import PyPDF2
import docx
import re # Added for robust JSON extraction
import yaml # For parsing template front-matter
import asyncio
import logging
import sys

# Configure logging to display info level messages
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

from pydantic import BaseModel
# Import specialized agent modules
import agent.Law as law_agents
from agent.router import classifier_agent
from agent.templatizer import templatizer_agent # Import the new agent
from agent.prefiller import prefiller_agent # Import the prefiller agent
from agent.question_generator import question_generator_agent # Import the question generator agent
from agent.bootstrap_agent import bootstrap_agent
from config import settings


from prisma import Prisma

# --- Database Client ---
# Instantiate the Prisma client
# CREATED BY "UOIONHHC"
db = Prisma()

# --- Pydantic Models for API Payloads ---

class TemplateIn(BaseModel):
    template_markdown: str

class DraftRequest(BaseModel):
    query: str

class FillTemplateRequest(BaseModel):
    template_id: str
    variables: Dict[str, str]

class PrefillRequest(BaseModel):
    template_id: str
    query: str

class GenerateQuestionsRequest(BaseModel):
    template_id: str
    filled_variables: Dict[str, str]




async def process_document_in_background(document_id: str, file_path: str, file_extension: str):
    """
    This function runs in the background to process the document.
    It extracts text, generates insights, and updates the database.
    """
    logging.info(f"Starting background processing for document: {document_id}")
    try:
        await db.document.update(where={"id": document_id}, data={"status": "processing"})

        # 1. Extract text from the document
        text_content = extract_text_from_file(file_path, file_extension)
        if not text_content:
            logging.error(f"Could not extract text from document {document_id}")
            await db.document.update(where={"id": document_id}, data={"status": "failed"})
            return

        # 2. Classify the document to select the correct agents
        logging.info(f"Classifying document {document_id}...")
        doc_type_raw = await classifier_agent.process(text_content)
        doc_type = doc_type_raw.strip().lower() # e.g., "banking", "legal"
        logging.info(f"Document {document_id} classified as: {doc_type}")

        # Select the appropriate agent set based on classification
        # Default to legal agents. This can be expanded if more domains are added.
        summarizer = law_agents.summarizer_agent
        entity_extractor = law_agents.entity_extractor_agent

        # 3. Use selected agents to generate insights asynchronously
        summary_task = summarizer.process(text_content)
        entities_task = entity_extractor.process(text_content)
        summary, entities_raw = await asyncio.gather(summary_task, entities_task)

        logging.info(f"Insights generated for document: {document_id}")
        
        # Robustly extract JSON from the model's raw output
        try:
            # Find the JSON block within the raw string, which might be wrapped in markdown
            json_match = re.search(r"```(json)?\s*(\{.*?\})\s*```", entities_raw, re.DOTALL)
            if json_match:
                json_str = json_match.group(2)
            else:
                json_str = entities_raw # Fallback to assuming the whole string is JSON
            entities = json.loads(json_str)
        except (json.JSONDecodeError, AttributeError):
            entities = {"error": "Failed to parse entities from model output. The format was invalid.", "raw_output": entities_raw}

        # 4. Store the final results in the database
        await db.document.update(
            where={"id": document_id},
            data={
                "status": "completed",
                "insights": json.dumps({"summary": summary, "entities": entities}),
                "documentType": doc_type,
                "fullText": text_content,
            },
        )
        logging.info(f"Successfully processed document: {document_id}")
        logging.info(f"Output for {document_id}: \nSUMMARY: {summary}\nENTITIES: {entities}")

    except Exception as e:
        logging.error(f"Error processing document {document_id}: {e}", exc_info=True)
        await db.document.update(where={"id": document_id}, data={"status": "failed"})

# --- Helper function for text extraction ---
def extract_text_from_file(file_path: str, file_extension: str) -> str:
    text = ""
    if file_extension == ".pdf":
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() or ""
    elif file_extension == ".docx":
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    elif file_extension == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
    return text

app = FastAPI(
    title="Intelligent Document Analysis Agent",
    description="API for AI-powered legal document review and insight generation.",
    version="0.2.0",
)

# Configure CORS for frontend interaction
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], 
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Directory to store uploaded documents temporarily
UPLOAD_DIR = "uploaded_documents"
os.makedirs(UPLOAD_DIR, exist_ok=True)

@app.on_event("startup")
async def startup():
    try:
        logging.info("Attempting to connect to the database...")
        await db.connect()
        logging.info("Database connection successful.")
    except Exception as e:
        logging.error("--- DATABASE CONNECTION FAILED ---")
        logging.error("Could not connect to the database. Please check the following:")
        logging.error("1. Your `DATABASE_URL` in the .env file is correct.")
        logging.error("2. The database server is running and accessible (e.g., not paused on Supabase).")
        logging.error("3. Your network/firewall is not blocking the connection on port 5432.")
        sys.exit(1) # Exit the application if the database is not available.

@app.on_event("shutdown")
async def shutdown():
    await db.disconnect()


# @app.post("/analyze")
# async def analyze_document(file: UploadFile):
#     text = (await file.read()).decode("utf-8")
#     result = analyze_legal_document(text)
#     return JSONResponse(content=result)

@app.post("/upload-document/")
async def upload_document(background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    """
    Uploads a document for intelligent analysis.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided.")

    file_extension = os.path.splitext(file.filename)[1].lower()
    if file_extension not in [".pdf", ".docx", ".txt"]: # Simplified for now, excluding images
        raise HTTPException(status_code=400, detail="Unsupported file type. Please upload PDF, DOCX, TXT, JPG, or PNG.")

    # Create a preliminary record in the database
    new_doc = await db.document.create(data={"status": "uploading"})
    document_id = new_doc.id
    file_path = os.path.join(UPLOAD_DIR, f"{document_id}{file_extension}")

    try:
        with open(file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        logging.info(f"File '{file.filename}' uploaded. Saved to: {file_path}")
        # Add the long-running processing task to the background
        background_tasks.add_task(
            process_document_in_background, document_id, file_path, file_extension
        )
        logging.info(f"Scheduled background processing for document_id: {document_id}")
        
        return {"message": "Document upload successful. Processing has started in the background.", "document_id": document_id}
    except Exception as e:
        # If anything fails, update the status in the DB
        await db.document.update(where={"id": document_id}, data={"status": "failed"})
        raise HTTPException(status_code=500, detail=f"Failed to upload or initiate processing: {str(e)}")

@app.post("/create-template-from-upload/")
async def create_template_from_upload(file: UploadFile = File(...)):
    """
    Uploads a document (.docx, .pdf) and uses an LLM to convert it into a
    reusable Markdown template with a YAML front-matter for variables.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided.")

    file_extension = os.path.splitext(file.filename)[1].lower()
    if file_extension not in [".pdf", ".docx", ".txt"]:
        raise HTTPException(status_code=400, detail="Unsupported file type. Please upload PDF, DOCX, or TXT.")

    # Save file temporarily to extract text
    file_path = os.path.join(UPLOAD_DIR, f"temp_{file.filename}")
    with open(file_path, "wb") as buffer:
        buffer.write(await file.read())

    text_content = extract_text_from_file(file_path, file_extension)
    os.remove(file_path) # Clean up the temporary file

    template_markdown = await templatizer_agent.process(text_content)
    return {"message": "Template extraction successful. Review the generated template.", "template_markdown": template_markdown}
@app.post("/save-template/")
async def save_template(template_data: TemplateIn):
    """
    Parses a markdown string with YAML front-matter and saves it as a
    template and its associated variables in the database.
    """
    try:
        # Split front-matter from body
        parts = template_data.template_markdown.split('---')
        if len(parts) < 3:
            raise HTTPException(status_code=400, detail="Invalid template format. Missing YAML front-matter.")

        front_matter_str = parts[1]
        body_md = "---".join(parts[2:]).strip()

        # Parse YAML front-matter
        try:
            parsed_yaml = yaml.safe_load(front_matter_str)
        except yaml.YAMLError as e:
            raise HTTPException(status_code=400, detail=f"Error parsing YAML front-matter: {e}")

        # Prepare variables with default keys if missing
        variables_to_create = []
        for idx, var in enumerate(parsed_yaml.get("variables", [])):
            if not isinstance(var, dict):
                continue
            variables_to_create.append({
                "key": var.get("key") or f"field_{idx}",
                "label": var.get("label") or f"Field {idx}",
                "description": var.get("description", ""),
                "example": var.get("example", ""),
                "required": var.get("required", True),
            })

        # Create template with variables in a single transaction
        new_template = await db.template.create(
            data={
                "title": parsed_yaml.get("title", "Untitled Template"),
                "fileDescription": parsed_yaml.get("file_description", ""),
                "jurisdiction": parsed_yaml.get("jurisdiction", ""),
                "docType": parsed_yaml.get("doc_type", ""),
                "similarityTags": parsed_yaml.get("similarity_tags", []),
                "bodyMd": body_md,
                "variables": {"create": variables_to_create} if variables_to_create else None
            }
        )

        return {"message": "Template saved successfully!", "template_id": new_template.id}

    except Exception as e:
        logging.error(f"Error saving template: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"An unexpected error occurred while saving the template: {str(e)}")

@app.get("/templates/", response_model=List[TemplateOut])
async def get_all_templates():
    templates = await db.template.find_many(
        order={"createdAt": "desc"},
        include={"variables": True}
    )
    return [TemplateOut.from_orm(t) for t in templates]
@app.get("/documents/")
async def get_all_documents():
    """
    Retrieves all documents with their details.
    """
    try:
        docs = await db.document.find_many()
        if not docs:
            return {"message": "No documents found."}
        return {"documents": docs}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error fetching documents: {str(e)}")

@app.get("/document/{document_id}/status")
async def get_processing_status(document_id: str):
    """
    Checks the processing status of a document.
    """
    doc = await db.document.find_unique(where={"id": document_id})
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found.")
    return {"document_id": document_id, "status": doc.status}

@app.get("/document/{document_id}/insights")
async def get_document_insights_endpoint(document_id: str) -> Dict:
    """
    Retrieves the analysis insights for a processed document.
    """
    doc = await db.document.find_unique(where={"id": document_id})
    if not doc or doc.status != "completed":
        raise HTTPException(status_code=404, detail="Insights not found or document not processed.")
    return doc.insights or {}

@app.post("/document/{document_id}/query")
async def query_document_endpoint(document_id: str, question: Dict[str, str]):
    """
    Allows users to ask specific questions about a processed document.
    """
    doc = await db.document.find_unique(where={"id": document_id})
    if not doc or not doc.fullText:
        raise HTTPException(status_code=404, detail="Document not found or not processed.")

    user_question = question.get("question")
    if not user_question:
        raise HTTPException(status_code=400, detail="Question not provided.")
    
    # Select the correct QA agent based on the stored document type
    doc_type = doc.documentType
    qa_agent_instance = law_agents.qa_agent # Default to legal QA agent

    # Use the agent to answer the question based on the stored full text
    answer = await qa_agent_instance.process(user_question, context=doc.fullText)
    return {"document_id": document_id, "question": user_question, "answer": answer}
@app.post("/find-templates")
async def find_templates(request: DraftRequest):
    """
    Finds relevant templates based on a user's query.
    Falls back to Web Bootstrap if no matches are found.
    """
    query_words = set(request.query.lower().split())
    try:
        templates = await db.template.find_many(include={"variables": True})
        scored_templates = []

        for t in templates:
            score = 0
            searchable_text = " ".join(filter(None, [
                t.title,
                t.fileDescription,
                t.docType,
                t.jurisdiction,
                " ".join(t.similarityTags)
            ])).lower()
            for word in query_words:
                if word in searchable_text:
                    score += 1
            if score > 0:
                scored_templates.append({"template": t, "score": score})

        # If found, return normally
        if scored_templates:
            scored_templates.sort(key=lambda x: x["score"], reverse=True)
            return {"status": "found", "results": scored_templates}

        # Otherwise → try Web Bootstrap
        logging.info(f"No local template match for query: {request.query}")
        new_template_data = await bootstrap_agent.bootstrap_template(request.query)
        
        if not new_template_data:
            return {"status": "not_found", "message": "No templates found online or locally."}
        
        # Defensive extraction of markdown content
        template_markdown = (
            new_template_data.get("template_markdown") or 
            new_template_data.get("full_markdown")
        )
        
        if not template_markdown:
            logging.error(f"No markdown content in bootstrap response. Keys: {new_template_data.keys()}")
            return {"status": "not_found", "message": "Bootstrap returned invalid template data."}
        
        # Parse YAML front-matter and body
        if template_markdown.strip().startswith("---"):
            parts = template_markdown.split("---", 2)  # Split into max 3 parts
            
            if len(parts) < 3:
                logging.error(f"Invalid YAML format. Parts count: {len(parts)}")
                return {"status": "error", "message": "Invalid template format from bootstrap."}
            
            yaml_part = parts[1].strip()
            body_md = parts[2].strip()
        else:
            logging.warning("Template doesn't start with YAML front-matter, using raw content")
            yaml_part = ""
            body_md = template_markdown

        # Parse YAML metadata
        import yaml
        try:
            meta = yaml.safe_load(yaml_part) if yaml_part else {}
        except yaml.YAMLError as e:
            logging.error(f"YAML parsing error: {e}")
            meta = {}

        # Extract variables with defaults
        variables_data = meta.get("variables", [])
        if not isinstance(variables_data, list):
            logging.warning(f"Variables is not a list: {type(variables_data)}")
            variables_data = []

        # Create template in database
        try:
            new_template = await db.template.create(
                data={
                    "title": meta.get("title") or new_template_data.get("title") or f"Template for {request.query}",
                    "fileDescription": meta.get("file_description") or new_template_data.get("file_description") or "",
                    "jurisdiction": meta.get("jurisdiction") or new_template_data.get("jurisdiction") or "",
                    "docType": meta.get("doc_type") or new_template_data.get("doc_type") or "",
                    "similarityTags": meta.get("similarity_tags") or new_template_data.get("similarity_tags") or [],
                    "bodyMd": body_md,
                    "variables": {
                        "create": [
                            {
                                "key": v.get("key", f"field_{idx}"),
                                "label": v.get("label", f"Field {idx}"),
                                "description": v.get("description", ""),
                                "example": v.get("example", ""),
                                "required": v.get("required", True),
                                "type": v.get("type", "string"),
                            }
                            for idx, v in enumerate(variables_data)
                            if isinstance(v, dict) and v.get("key")  # Only valid variable objects
                        ]
                    }
                }
            )

            return {
                "status": "bootstrapped",
                "source_url": new_template_data.get("source_url"),
                "source_title": new_template_data.get("source_title"),
                "template": new_template
            }
        
        except Exception as db_error:
            logging.error(f"Database error creating template: {db_error}", exc_info=True)
            raise HTTPException(status_code=500, detail=f"Failed to save template: {str(db_error)}")

    except HTTPException:
        # Re-raise HTTP exceptions
        raise
    except Exception as e:
        logging.error(f"Error in find_templates: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error searching or bootstrapping templates: {str(e)}")
@app.post("/fill-template")
async def fill_template(request: FillTemplateRequest):
    """
    Fills a template with provided variables to generate a document draft.
    """
    template = await db.template.find_unique(where={"id": request.template_id})
    if not template:
        raise HTTPException(status_code=404, detail="Template not found.")

    draft_content = template.bodyMd
    for key, value in request.variables.items():
        # Simple string replacement for placeholders like {{variable_key}}
        draft_content = draft_content.replace(f"{{{{{key}}}}}", str(value))

    return {"draft_markdown": draft_content}

@app.post("/prefill-variables-from-query")
async def prefill_variables_from_query(request: PrefillRequest):
    """
    
    Uses an LLM to pre-fill template variables based on the initial user query.
    """
    template = await db.template.find_unique(where={"id": request.template_id}, include={"variables": True})
    if not template:
        raise HTTPException(status_code=404, detail="Template not found.")

    # Prepare the variables for the agent prompt
    variables_for_prompt = [
        {"key": v.key, "label": v.label, "description": v.description}
        for v in template.variables
    ]

    try:
        detected_variables = await prefiller_agent.process({
            "query": request.query,
            "variables_json": json.dumps(variables_for_prompt, indent=2)
        })

        return {"message": "Prefill successful.", "query": request.query, "detected_variables": detected_variables}

    except Exception as e:
        logging.error(f"Error during prefill for template {request.template_id}: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to prefill variables from query.")

@app.post("/generate-questions-for-missing-variables")
async def generate_questions(request: GenerateQuestionsRequest):
    """
    Generates human-friendly questions for required variables that are not yet filled.
    """
    template = await db.template.find_unique(where={"id": request.template_id}, include={"variables": True})
    if not template:
        raise HTTPException(status_code=404, detail="Template not found.")

    missing_vars = [
        var for var in template.variables 
        if var.required and var.key not in request.filled_variables
    ]

    # Create a list of tasks to generate questions concurrently
    question_generation_tasks = []
    for var in template.variables:
        if var.required and var.key not in request.filled_variables:
            task = question_generator_agent.process({
                "label": var.label,
                "description": var.description
            })
            question_generation_tasks.append(task)

    # Execute all question generation tasks in parallel
    generated_questions = await asyncio.gather(*question_generation_tasks)

    # Combine the results with the variable keys
    questions_to_ask = [
        {"key": var.key, "question": question, "example": var.example}
        for var, question in zip(missing_vars, generated_questions)
    ]

    return {"missing_variables_questions": questions_to_ask}

# CREATED BY "UOIONHHC"
@app.post("/export/")
async def export_document(
    body: str = Form(...),
    filename: str = Form(...),
    filetype: str = Form(...),
):
    """
    Exports the given markdown body to a .docx or .pdf file.
    """
    if filetype not in ["docx", "pdf"]:
        raise HTTPException(status_code=400, detail="Unsupported filetype. Use 'docx' or 'pdf'.")

    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{filetype}") as tmp_file:
        if filetype == "docx":
            from docx import Document
            document = Document()
            # A simple approach to add markdown as paragraphs.
            # For more complex markdown, a more robust parser would be needed.
            for line in body.split('\n'):
                document.add_paragraph(line)
            document.save(tmp_file.name)
        elif filetype == "pdf":
            from md2pdf.core import md2pdf
            # Using md2pdf which in turn uses pandoc.
            # Ensure pandoc is installed on your system.
            try:
                md2pdf(tmp_file.name, md_content=body)
            except Exception as e:
                logging.error(f"PDF conversion failed: {e}")
                raise HTTPException(status_code=500, detail="PDF conversion failed. Ensure pandoc is installed.")

        return FileResponse(
            path=tmp_file.name,
            filename=f"{filename}.{filetype}",
            media_type=f"application/vnd.openxmlformats-officedocument.wordprocessingml.document" if filetype == "docx" else "application/pdf",
        )
